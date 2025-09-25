"""
Test Configuration and Utilities
AI-Powered CV Processing Engine Test Suite

This module provides configuration, fixtures, and utilities for testing
the CV processing system.

Version: 1.0
Date: September 24, 2025
"""

import os
import tempfile
import base64
from typing import Dict, Any
from docx import Document
from docx.shared import Inches
import io
import logging

# Configure test logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

class TestDataGenerator:
    """Generate test data for CV processing tests"""
    
    @staticmethod
    def create_sample_pdf_bytes() -> bytes:
        """Create sample PDF file bytes for testing"""
        # For testing purposes, we'll use placeholder bytes
        # In a real scenario, you might use a PDF library to create actual PDF content
        pdf_header = b"%PDF-1.4\n"
        pdf_content = b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        pdf_content += b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        pdf_content += b"3 0 obj\n<< /Type /Page /Parent 2 0 R /Contents 4 0 R >>\nendobj\n"
        pdf_content += b"4 0 obj\n<< /Length 44 >>\nstream\nBT\n/F1 12 Tf\n100 100 Td\n(Test CV Content) Tj\nET\nendstream\nendobj\n"
        pdf_content += b"xref\n0 5\n0000000000 65535 f \n0000000010 00000 n \n0000000053 00000 n \n0000000125 00000 n \n0000000185 00000 n \ntrailer\n<< /Size 5 /Root 1 0 R >>\nstartxref\n285\n%%EOF"
        
        return pdf_header + pdf_content
    
    @staticmethod
    def create_sample_docx_bytes() -> bytes:
        """Create sample DOCX file bytes for testing"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('John Smith', 0)
        doc.add_heading('Senior Software Engineer', level=1)
        
        # Add contact information
        doc.add_paragraph('Email: john.smith@email.com')
        doc.add_paragraph('Phone: +1-555-123-4567')
        doc.add_paragraph('LinkedIn: https://linkedin.com/in/johnsmith')
        doc.add_paragraph('GitHub: https://github.com/johnsmith')
        
        # Add professional summary
        doc.add_heading('Professional Summary', level=2)
        doc.add_paragraph(
            'Experienced software engineer with 8 years of experience in full-stack development. '
            'Expert in Python, JavaScript, and cloud architecture.'
        )
        
        # Add work experience
        doc.add_heading('Work Experience', level=2)
        doc.add_heading('Senior Software Engineer | Tech Innovations Inc | 2020-Present', level=3)
        doc.add_paragraph('• Led development team of 5 engineers')
        doc.add_paragraph('• Architected microservices solutions using AWS')
        doc.add_paragraph('• Improved system performance by 40%')
        
        # Add education
        doc.add_heading('Education', level=2)
        doc.add_paragraph('Master of Computer Science | Stanford University | 2018')
        doc.add_paragraph('Bachelor of Computer Science | UC Berkeley | 2016')
        
        # Add skills
        doc.add_heading('Skills', level=2)
        doc.add_paragraph('Programming Languages: Python, JavaScript, Java, Go')
        doc.add_paragraph('Frameworks: Django, React, Node.js')
        doc.add_paragraph('Cloud Platforms: AWS, Azure')
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.read()
    
    @staticmethod
    def create_multilingual_cv_content() -> str:
        """Create CV content with multilingual elements for testing"""
        return """
# محمد أحمد (Mohammed Ahmed)
## Software Engineer | مهندس برمجيات

**Email:** mohammed.ahmed@email.com  
**Phone:** +971-50-123-4567  
**LinkedIn:** https://linkedin.com/in/mohammedahmed  

### Current Position
Senior Software Engineer at Emirates Tech Solutions  
Location: Dubai, UAE | دبي، الإمارات العربية المتحدة  

### Professional Summary
Experienced software engineer with 6 years of experience.
خبرة 6 سنوات في تطوير البرمجيات

### Education
**Bachelor of Computer Science** | American University of Sharjah | 2018  
**بكالوريوس علوم الحاسوب** | الجامعة الأمريكية في الشارقة  

### Certifications  
- Emirates ID: 784-1985-1234567-8
- AWS Certified Solutions Architect
- Microsoft Azure Certified

### Skills
- Programming: Python, JavaScript, C#
- Languages: Arabic (Native), English (Fluent)
- اللغات: العربية (الأم)، الإنجليزية (بطلاقة)
"""
    
    @staticmethod
    def create_complex_pii_cv() -> str:
        """Create CV with complex PII scenarios for testing"""
        return """
# Dr. John Michael Smith Jr.
## Chief Technology Officer

**Personal Information:**
- Full Name: John Michael Smith Jr.
- Email: j.smith@techcorp.com, john.smith.personal@gmail.com
- Phone: +1-555-123-4567, (555) 987-6543
- Address: 123 Tech Street, San Francisco, CA 94105
- SSN: 123-45-6789
- Driver's License: D123456789

**Professional Profiles:**
- LinkedIn: https://linkedin.com/in/johnmichaelsmith
- GitHub: https://github.com/jmsmith
- Personal Website: https://johnmichaelsmith.com
- Twitter: @johnmsmith

**Current Employment:**
Chief Technology Officer at TechCorp Inc.
Previous: Senior Director at InnovateTech Solutions
Previous: Lead Engineer at StartupXYZ

**Education:**
- Ph.D. Computer Science, Stanford University (2015)
- M.S. Computer Engineering, MIT (2012)  
- B.S. Computer Science, UC Berkeley (2010)

**References:**
- Dr. Jane Doe, Professor at Stanford University (jane.doe@stanford.edu)
- Mr. Bob Johnson, CEO at TechCorp Inc. (bob.johnson@techcorp.com)

**Certifications:**
- PMP Certificate Number: 1234567
- AWS Certification ID: AWS-12345-67890
- Security Clearance: Secret (Facility: NSA-SF-001)
"""
    
    @staticmethod
    def create_edge_case_content() -> str:
        """Create content with edge cases for testing"""
        return """
# Test User
## Position with Special Characters & Symbols!

Email: test@domain-with-hyphens.co.uk
Phone: +44-207-123-4567 (UK), +1-555-987-6543 (US)
Website: https://test-user.github.io/portfolio

**Experience with Numbers:**
- 5+ years experience
- Managed $2.5M budget  
- Led team of 15+ developers
- Achieved 99.9% uptime

**Special Characters in Text:**
- Skills: C++, C#, .NET, Node.js, Vue.js
- Frameworks: Angular 2+, React 16+
- Databases: MySQL 8.0, PostgreSQL 13+

**URLs and Links:**
- https://subdomain.example.com/path?param=value&other=123
- ftp://files.example.com/documents/
- mailto:contact@example.com

**Mixed Formats:**
- Date: 01/15/2023, 2023-01-15, January 15, 2023
- Currency: $50,000, €45,000, £40,000
- Percentages: 95%, 12.5%, 0.01%
"""

class TestEnvironmentSetup:
    """Setup and teardown utilities for test environment"""
    
    @staticmethod
    def setup_test_environment():
        """Setup test environment variables"""
        test_env = {
            'AZURE_OPENAI_ENDPOINT': 'https://test-openai.openai.azure.com/',
            'AZURE_OPENAI_API_KEY': 'test-api-key-12345',
            'AZURE_OPENAI_API_VERSION': '2024-02-15-preview',
            'CV_DATA_EXTRACTOR_ASSISTANT_ID': 'asst_cv_data_extractor_test',
            'CV_PII_IDENTIFIER_ASSISTANT_ID': 'asst_pii_identifier_test',
            'CV_SKILLS_ANALYST_ASSISTANT_ID': 'asst_skills_analyst_test'
        }
        
        for key, value in test_env.items():
            os.environ[key] = value
        
        return test_env
    
    @staticmethod
    def cleanup_test_environment():
        """Clean up test environment variables"""
        test_vars = [
            'AZURE_OPENAI_ENDPOINT',
            'AZURE_OPENAI_API_KEY',
            'AZURE_OPENAI_API_VERSION',
            'CV_DATA_EXTRACTOR_ASSISTANT_ID',
            'CV_PII_IDENTIFIER_ASSISTANT_ID',
            'CV_SKILLS_ANALYST_ASSISTANT_ID'
        ]
        
        for var in test_vars:
            if var in os.environ:
                del os.environ[var]
    
    @staticmethod
    def create_temp_file(content: bytes, suffix: str = '.tmp') -> str:
        """Create a temporary file with given content"""
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(content)
            return temp_file.name

class MockResponseGenerator:
    """Generate mock responses for AI services"""
    
    @staticmethod
    def create_candidate_extraction_response() -> Dict[str, Any]:
        """Create mock response for candidate data extraction"""
        return {
            "candidateProfile": {
                "candidateName": "John Smith",
                "jobTitle": "Senior Software Engineer",
                "status": "Active",
                "candidateEmail": "john.smith@email.com",
                "candidatePhone": "+1-555-123-4567",
                "linkedInUrl": "https://linkedin.com/in/johnsmith",
                "otherProfileUrls": ["https://github.com/johnsmith", "https://johnsmith.dev"],
                "currentLocation": "New York, NY",
                "currentTitle": "Senior Software Engineer",
                "currentCompany": "Tech Innovations Inc",
                "professionalHeadline": "Experienced Software Engineer with Cloud Architecture Expertise",
                "totalExperienceYears": 8,
                "relevantExperienceYears": 6,
                "highestQualification": "Master of Computer Science",
                "currentSalary": "120000 USD",
                "expectedSalary": "140000 USD",
                "availabilityStatus": "Available in 30 days",
                "visaStatus": "Permanent Resident"
            }
        }
    
    @staticmethod
    def create_pii_identification_response() -> Dict[str, Any]:
        """Create mock response for PII identification"""
        return {
            "pii_entities": [
                {
                    "original_value": "John Smith",
                    "pii_type": "name",
                    "sensitivity_level": 4,
                    "all_variations": ["John Smith", "John", "Smith"],
                    "context": "candidate name in header",
                    "region_specific": False
                },
                {
                    "original_value": "Tech Innovations Inc",
                    "pii_type": "organization",
                    "sensitivity_level": 2,
                    "all_variations": ["Tech Innovations Inc"],
                    "context": "current employer",
                    "region_specific": False
                },
                {
                    "original_value": "Stanford University",
                    "pii_type": "school",
                    "sensitivity_level": 2,
                    "all_variations": ["Stanford University", "Stanford"],
                    "context": "education institution",
                    "region_specific": False
                }
            ]
        }
    
    @staticmethod
    def create_skills_analysis_response() -> Dict[str, Any]:
        """Create mock response for skills analysis"""
        return {
            "analysisMetrics": {
                "aiConfidenceScore": 85,
                "aiRemarks": "Strong technical background with relevant industry experience. Excellent leadership skills demonstrated through team management. Good fit for senior technical roles.",
                "overallFitScore": 78
            },
            "professionalProfile": {
                "coreSkills": [
                    {
                        "skillName": "Python",
                        "proficiencyLevel": "Expert",
                        "yearsOfExperience": 5
                    },
                    {
                        "skillName": "Cloud Architecture",
                        "proficiencyLevel": "Advanced",
                        "yearsOfExperience": 4
                    },
                    {
                        "skillName": "Team Leadership",
                        "proficiencyLevel": "Expert",
                        "yearsOfExperience": 3
                    }
                ],
                "workExperience": [
                    {
                        "jobTitle": "Senior Software Engineer",
                        "companySize": "Large Enterprise",
                        "duration": "3 years 2 months",
                        "keyResponsibilities": [
                            "Led development team of 5 engineers",
                            "Architected microservices solutions",
                            "Mentored junior developers"
                        ],
                        "keyAchievements": [
                            "Improved system performance by 40%",
                            "Reduced deployment time by 60%",
                            "Successfully delivered 15+ projects"
                        ]
                    }
                ],
                "educationProfile": {
                    "highestDegreeLevel": "Masters",
                    "fieldOfStudy": "Computer Science",
                    "relevantCertifications": [
                        "AWS Certified Solutions Architect",
                        "PMP Certified"
                    ]
                }
            }
        }

class ValidationHelpers:
    """Helper functions for test validations"""
    
    @staticmethod
    def validate_candidate_profile_structure(profile: Dict[str, Any]) -> bool:
        """Validate that candidate profile has all required fields"""
        required_fields = [
            'candidateName', 'jobTitle', 'status', 'candidateEmail',
            'candidatePhone', 'linkedInUrl', 'otherProfileUrls',
            'currentLocation', 'currentTitle', 'currentCompany',
            'professionalHeadline', 'totalExperienceYears', 'relevantExperienceYears',
            'highestQualification', 'currentSalary', 'expectedSalary',
            'availabilityStatus', 'visaStatus'
        ]
        
        for field in required_fields:
            if field not in profile:
                return False
        
        return True
    
    @staticmethod
    def validate_pii_entity_structure(entity: Dict[str, Any]) -> bool:
        """Validate that PII entity has all required fields"""
        required_fields = [
            'original_value', 'pii_type', 'sensitivity_level',
            'all_variations', 'context'
        ]
        
        for field in required_fields:
            if field not in entity:
                return False
        
        # Validate sensitivity level is in valid range
        if not (1 <= entity['sensitivity_level'] <= 4):
            return False
        
        # Validate PII type is valid
        valid_types = [
            'name', 'address', 'organization', 'school', 'job_title',
            'certification_number', 'license_number', 'national_id'
        ]
        if entity['pii_type'] not in valid_types:
            return False
        
        return True
    
    @staticmethod
    def validate_skills_analysis_structure(analysis: Dict[str, Any]) -> bool:
        """Validate that skills analysis has all required fields"""
        if 'analysisMetrics' not in analysis or 'professionalProfile' not in analysis:
            return False
        
        # Validate analysis metrics
        metrics_fields = ['aiConfidenceScore', 'aiRemarks', 'overallFitScore']
        for field in metrics_fields:
            if field not in analysis['analysisMetrics']:
                return False
        
        # Validate professional profile
        profile_fields = ['coreSkills', 'workExperience', 'educationProfile']
        for field in profile_fields:
            if field not in analysis['professionalProfile']:
                return False
        
        return True
    
    @staticmethod
    def validate_ai_call_log_structure(log: Dict[str, Any]) -> bool:
        """Validate that AI call log has all required fields"""
        required_fields = [
            'stage', 'assistantId', 'modelUsed', 'inputPrompt',
            'inputTokens', 'responseText', 'outputTokens', 
            'callDurationMs', 'callTimestamp', 'success'
        ]
        
        for field in required_fields:
            if field not in log:
                return False
        
        # Validate stage is valid
        valid_stages = ['candidate_data_extraction', 'pii_identification', 'skills_analysis']
        if log['stage'] not in valid_stages:
            return False
        
        return True
